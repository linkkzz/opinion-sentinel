# -*- coding: utf-8 -*-
"""Markdown -> docx 转换脚本（针对本技术方案文档定制）。"""
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+?)`")


def set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    tc_pr.append(shd)


def add_runs_with_inline(paragraph, text: str):
    """解析行内 **加粗** / `代码` 并写入段落。"""
    # 先按 ** 拆出加粗段
    pos = 0
    for m in INLINE_BOLD.finditer(text):
        if m.start() > pos:
            _add_inline_seg(paragraph, text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        _add_inline_seg(paragraph, text[pos:])


def _add_inline_seg(paragraph, seg: str):
    """处理段内 `代码` 与普通文本。"""
    pos = 0
    for m in INLINE_CODE.finditer(seg):
        if m.start() > pos:
            paragraph.add_run(seg[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10)
        pos = m.end()
    if pos < len(seg):
        paragraph.add_run(seg[pos:])


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = parse_table_row(s)
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells) and len(cells) > 0


def main(md_path: str, docx_path: str):
    doc = Document()

    # 全局默认字体（正文）
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 页边距
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for cl in code_lines:
                run = p.add_run(cl + "\n" if cl else "\n")
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            run = h.add_run(text)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if level == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            elif level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)
            else:
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # 分隔线
        if stripped == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pbdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "6",
                qn("w:space"): "1",
                qn("w:color"): "BFBFBF",
            })
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        # 图片 ![caption](path)
        m = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", stripped)
        if m:
            caption = m.group(1).strip()
            img_path = m.group(2).strip()
            # 相对路径基于 md 文件目录解析
            if not os.path.isabs(img_path):
                img_path = os.path.join(os.path.dirname(os.path.abspath(md_path)), img_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            # 宽度自适应：限制不超过页宽(约15cm)
            max_w = Inches(5.9)
            if os.path.exists(img_path):
                # 读取图片真实宽高比
                try:
                    from PIL import Image
                    with Image.open(img_path) as im:
                        w, h = im.size
                    aspect = h / w
                    run.add_picture(img_path, width=max_w)
                except Exception:
                    # 无 PIL 则直接按宽度插入
                    run.add_picture(img_path, width=max_w)
            else:
                run.add_text(f"[图片缺失: {m.group(2)}]")
            # 图注
            if caption:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(10)
                crun = cap.add_run(caption)
                crun.font.size = Pt(9)
                crun.italic = True
                crun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 表格
        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            # 去掉分隔行
            data_rows = [r for r in rows if not is_table_separator(r)]
            if not data_rows:
                continue
            parsed = [parse_table_row(r) for r in data_rows]
            ncols = max(len(r) for r in parsed)
            table = doc.add_table(rows=len(parsed), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row_cells in enumerate(parsed):
                for ci in range(ncols):
                    cell = table.cell(ri, ci)
                    text = row_cells[ci] if ci < len(row_cells) else ""
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    add_runs_with_inline(p, text)
                    if ri == 0:
                        set_cell_bg(cell, "1F497D")
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.bold = True
                            run.font.size = Pt(10)
                    else:
                        for run in p.runs:
                            run.font.size = Pt(9.5)
                        if ri % 2 == 0:
                            set_cell_bg(cell, "F2F6FB")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # 无序列表
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3 + indent * 0.02)
            add_runs_with_inline(p, text)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs_with_inline(p, text)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            pPr = p._p.get_or_add_pPr()
            pbdr = pPr.makeelement(qn("w:pBdr"), {})
            left = pbdr.makeelement(qn("w:left"), {
                qn("w:val"): "single", qn("w:sz"): "18",
                qn("w:space"): "8", qn("w:color"): "2E5C8A",
            })
            pbdr.append(left)
            pPr.append(pbdr)
            add_runs_with_inline(p, text)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.4
        add_runs_with_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"OK: {docx_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
